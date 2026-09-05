"""Formatting-only tidy pass over the B13 report docx.

Reads writeup/local/report/B13_report.docx and writes
writeup/local/report/B13_report_tidy.docx.

Nothing textual is touched: no paragraph, run, table or image is added,
removed or re-worded.  Only styles, section geometry, table geometry,
inline-image extents and a page-number footer are changed.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "writeup" / "local" / "report" / "B13_report.docx"
DST = REPO / "writeup" / "local" / "report" / "B13_report_tidy.docx"

BODY_FONT = "Georgia"
MONO_FONT = "Consolas"
TEXT_WIDTH_CM = 17.0  # A4 (21 cm) minus 2 cm margins each side.
CAPTION_PREFIXES = ("Figure 1.", "Figure 2.", "Table ")


# --------------------------------------------------------------------------
# low-level xml helpers
# --------------------------------------------------------------------------
def _sub(parent, tag):
    """Get-or-create a direct child element."""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def set_style_font(style, *, name=None, size_pt=None, bold=None, italic=None):
    f = style.font
    if name is not None:
        f.name = name
        # Also pin the complex/east-asian slots so LibreOffice honours it.
        rpr = style.element.get_or_add_rPr()
        rfonts = _sub(rpr, "w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), name)
    if size_pt is not None:
        f.size = Pt(size_pt)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic


def set_paragraph_format(pf, *, before=None, after=None, line=None,
                         keep_with_next=None, keep_together=None):
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next
    if keep_together is not None:
        pf.keep_together = keep_together


# --------------------------------------------------------------------------
# page geometry
# --------------------------------------------------------------------------
def fix_sections(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        section.start_type = WD_SECTION_START.NEW_PAGE


# --------------------------------------------------------------------------
# styles
# --------------------------------------------------------------------------
def fix_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, name=BODY_FONT, size_pt=10.5)
    set_paragraph_format(normal.paragraph_format, before=0, after=6, line=1.15)

    for name in ("Body Text", "Default Paragraph Font", "Text Body"):
        try:
            st = styles[name]
        except KeyError:
            continue
        if st.type is not None and hasattr(st, "paragraph_format"):
            set_style_font(st, name=BODY_FONT, size_pt=10.5)
            set_paragraph_format(st.paragraph_format, before=0, after=6,
                                 line=1.15)

    heading_spec = {
        "Heading 1": (16.0, 16, 8),
        "Heading 2": (13.0, 14, 6),
        "Heading 3": (11.5, 12, 5),
        "Heading 4": (10.5, 10, 4),
        "Heading 5": (10.5, 9, 4),
        "Heading 6": (10.5, 9, 4),
    }
    for name, (size, before, after) in heading_spec.items():
        try:
            st = styles[name]
        except KeyError:
            continue
        set_style_font(st, name=BODY_FONT, size_pt=size, bold=True)
        set_paragraph_format(st.paragraph_format, before=before, after=after,
                             line=1.1, keep_with_next=True, keep_together=True)

    # Monospace / code character style produced by the HTML source.
    for name in ("Source Text", "Preformatted Text", "Teletype"):
        try:
            st = styles[name]
        except KeyError:
            continue
        set_style_font(st, name=MONO_FONT, size_pt=9.0)

    # Caption-ish paragraph styles emitted by the converter.
    for name in ("Body Text.figcaption", "Body Text.tabcaption", "Caption"):
        try:
            st = styles[name]
        except KeyError:
            continue
        set_style_font(st, name=BODY_FONT, size_pt=9.5, italic=True)
        set_paragraph_format(st.paragraph_format, before=3, after=8, line=1.1)

    # Figure container paragraphs: centre and keep tight to their caption.
    try:
        st = styles["Body Text.figure"]
    except KeyError:
        pass
    else:
        set_paragraph_format(st.paragraph_format, before=6, after=3, line=1.0,
                             keep_with_next=True)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        st = styles["Block Quotation"]
    except KeyError:
        pass
    else:
        set_style_font(st, name=BODY_FONT, size_pt=10.0)
        set_paragraph_format(st.paragraph_format, before=4, after=6, line=1.15)

    # Table styles used inside cells.
    for name, bold in (("Table Contents", None), ("Table Heading", True)):
        try:
            st = styles[name]
        except KeyError:
            continue
        set_style_font(st, name=BODY_FONT, size_pt=8.5, bold=bold)
        set_paragraph_format(st.paragraph_format, before=0, after=0, line=1.0)


# --------------------------------------------------------------------------
# captions detected by text prefix (some are plain Body Text paragraphs)
# --------------------------------------------------------------------------
def fix_captions(doc):
    n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text.startswith(CAPTION_PREFIXES):
            continue
        n += 1
        set_paragraph_format(para.paragraph_format, before=3, after=8,
                             line=1.1, keep_together=True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.name = BODY_FONT
    return n


# --------------------------------------------------------------------------
# tables
# --------------------------------------------------------------------------
def fix_tables(doc):
    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr

        # Full text width, autofit contents.
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        w = _sub(tbl_pr, "w:tblW")
        w.set(qn("w:w"), "5000")
        w.set(qn("w:type"), "pct")
        layout = _sub(tbl_pr, "w:tblLayout")
        layout.set(qn("w:type"), "autofit")

        # Minimal cell padding.
        margins = _sub(tbl_pr, "w:tblCellMar")
        for side, twips in (("top", 20), ("bottom", 20),
                            ("left", 60), ("right", 60)):
            m = _sub(margins, f"w:{side}")
            m.set(qn("w:w"), str(twips))
            m.set(qn("w:type"), "dxa")

        # Drop fixed per-cell/grid widths so autofit can act.
        for gridcol in tbl.findall(qn("w:tblGrid") + "/" + qn("w:gridCol")):
            pass  # grid hints are harmless once tblLayout is autofit

        for row in table.rows:
            row.allow_break_across_pages = True
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = _sub(tc_pr, "w:tcW")
                tc_w.set(qn("w:w"), "0")
                tc_w.set(qn("w:type"), "auto")
                for para in cell.paragraphs:
                    set_paragraph_format(para.paragraph_format, before=0,
                                         after=0, line=1.0)
                    for run in para.runs:
                        run.font.size = Pt(8.5)
        # Repeat header row on page breaks.
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))


# --------------------------------------------------------------------------
# inline images
# --------------------------------------------------------------------------
def fix_images(doc):
    target = Cm(TEXT_WIDTH_CM)
    out = []
    for shape in doc.inline_shapes:
        w, h = shape.width, shape.height
        if w and h:
            new_h = int(round(h * (target / float(w))))
            shape.width = target
            shape.height = new_h
            out.append((w, h, shape.width, shape.height))
    return out


# --------------------------------------------------------------------------
# footer with "Page X of Y"
# --------------------------------------------------------------------------
def _field_run(paragraph, instr):
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_txt = OxmlElement("w:t")
    fld_txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr_el, fld_sep, fld_txt, fld_end):
        r._r.append(el)
    r.font.name = BODY_FONT
    r.font.size = Pt(9)
    return r


def add_footer(doc):
    for section in doc.sections:
        section.different_first_page_header_footer = False
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.style = doc.styles["Normal"]
        set_paragraph_format(para.paragraph_format, before=0, after=0, line=1.0)

        def txt(s):
            run = para.add_run(s)
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            return run

        txt("page ")
        _field_run(para, " PAGE ")
        txt(" of ")
        _field_run(para, " NUMPAGES ")


# --------------------------------------------------------------------------
# verification
# --------------------------------------------------------------------------
def fingerprint(path):
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs)
    tbl_text = []
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                tbl_text.append(cell.text)
    return {
        "paragraphs": len(d.paragraphs),
        "tables": len(d.tables),
        "inline_shapes": len(d.inline_shapes),
        "text": text,
        "table_text": "\n".join(tbl_text),
    }


def main():
    doc = docx.Document(str(SRC))

    fix_sections(doc)
    fix_styles(doc)
    n_caps = fix_captions(doc)
    fix_tables(doc)
    imgs = fix_images(doc)
    add_footer(doc)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))

    before, after = fingerprint(SRC), fingerprint(DST)
    ok = True
    for key in ("paragraphs", "tables", "inline_shapes", "text", "table_text"):
        same = before[key] == after[key]
        ok &= same
        shown = before[key] if key not in ("text", "table_text") else f"{len(before[key])} chars"
        print(f"{key:>14}: {shown}  identical={same}")
    print(f"     captions: {n_caps} paragraphs restyled")
    for w, h, nw, nh in imgs:
        print(f"        image: {Cm(w).cm if False else w/360000:.2f}x{h/360000:.2f} cm "
              f"-> {nw/360000:.2f}x{nh/360000:.2f} cm")
    print(f"wrote {DST}")
    if not ok:
        print("TEXT CHANGED - aborting", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
